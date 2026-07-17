from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).parent
OUT = ROOT / "Karaoke_Booking_POS_Submission"
OUT.mkdir(exist_ok=True)


def save_use_case(path):
    im=Image.new('RGB',(1800,1200),'white'); d=ImageDraw.Draw(im)
    font=ImageFont.truetype('C:/Windows/Fonts/arial.ttf',22); small=ImageFont.truetype('C:/Windows/Fonts/arial.ttf',18); bold=ImageFont.truetype('C:/Windows/Fonts/arialbd.ttf',28)
    d.rounded_rectangle((340,60,1460,1140),radius=20,outline='#17365D',width=4); d.text((900,85),'KARAOKE BOOKING & POS',font=bold,fill='#17365D',anchor='mm')
    actors = {"Khách hàng": (130, 250), "Lễ tân/Thu ngân": (130, 770), "Nhân viên phục vụ": (1660, 250), "Quản lý chi nhánh": (1660, 700), "Quản trị hệ thống": (1660, 1030)}
    cases = {
        "Đăng ký / đăng nhập": (550, 240), "Tìm phòng & xem giá": (900, 240), "Đặt / hủy phòng": (1250, 330),
        "Check-in / chuyển phòng": (580, 520), "Gọi món / cập nhật order": (930, 510), "Thanh toán & xuất hóa đơn": (1240, 590),
        "Quản lý phòng & bảng giá": (580, 790), "Quản lý menu / tồn kho": (930, 800), "Báo cáo doanh thu": (1240, 790),
        "Quản lý người dùng & quyền": (900, 1030)
    }
    links = {
        "Khách hàng": ["Đăng ký / đăng nhập", "Tìm phòng & xem giá", "Đặt / hủy phòng", "Gọi món / cập nhật order"],
        "Lễ tân/Thu ngân": ["Đặt / hủy phòng", "Check-in / chuyển phòng", "Thanh toán & xuất hóa đơn"],
        "Nhân viên phục vụ": ["Gọi món / cập nhật order"],
        "Quản lý chi nhánh": ["Quản lý phòng & bảng giá", "Quản lý menu / tồn kho", "Báo cáo doanh thu"],
        "Quản trị hệ thống": ["Quản lý người dùng & quyền"]
    }
    for a,cs in links.items():
        for c in cs: d.line((actors[a],cases[c]),fill='#94A3B8',width=2)
    for name,(x,y) in actors.items():
        d.rounded_rectangle((x-125,y-35,x+125,y+35),radius=12,fill='#DBEAFE',outline='#2563EB',width=3); d.text((x,y),name,font=small,fill='#111827',anchor='mm')
    for name,(x,y) in cases.items():
        d.ellipse((x-145,y-42,x+145,y+42),fill='#F8FAFC',outline='#334155',width=3); d.text((x,y),name,font=small,fill='#111827',anchor='mm')
    im.save(path)


def save_erd(path):
    im=Image.new('RGB',(2200,1450),'white'); d=ImageDraw.Draw(im)
    font=ImageFont.truetype('C:/Windows/Fonts/arial.ttf',20); small=ImageFont.truetype('C:/Windows/Fonts/arial.ttf',17); bold=ImageFont.truetype('C:/Windows/Fonts/arialbd.ttf',24); title=ImageFont.truetype('C:/Windows/Fonts/arialbd.ttf',32)
    entities = {
        "BRANCH": (60, 180, ["PK branch_id", "name", "address", "status"]), "ROOM_TYPE": (520, 180, ["PK room_type_id", "name", "capacity", "base_rate"]),
        "ROOM": (980, 180, ["PK room_id", "FK branch_id", "FK room_type_id", "code", "status"]), "PRODUCT": (1710, 180, ["PK product_id", "FK branch_id", "name", "price, stock_qty"]),
        "CUSTOMER": (60, 650, ["PK customer_id", "phone (UQ)", "full_name", "loyalty_point"]), "BOOKING": (540, 650, ["PK booking_id", "FK customer_id", "FK room_id", "start_at, end_at", "status, deposit"]),
        "SESSION": (1010, 650, ["PK session_id", "FK booking_id?", "FK room_id", "check_in/out", "status"]), "ORDER": (1490, 650, ["PK order_id", "FK session_id", "FK staff_id", "status, total"]),
        "STAFF": (60, 1100, ["PK staff_id", "FK branch_id", "FK role_id", "username (UQ)"]), "PAYMENT": (540, 1100, ["PK payment_id", "FK invoice_id", "method, amount", "status, paid_at"]),
        "INVOICE": (1010, 1100, ["PK invoice_id", "FK session_id (UQ)", "subtotal, discount", "tax, grand_total"]), "ORDER_ITEM": (1710, 1100, ["PK order_item_id", "FK order_id", "FK product_id", "qty, unit_price"]),
    }
    boxes = {}
    for name, (x, y, fields) in entities.items():
        w,h=350,230; d.rounded_rectangle((x,y,x+w,y+h),radius=12,fill='#F8FAFC',outline='#0F172A',width=3); d.rectangle((x,y,x+w,y+45),fill='#BFDBFE',outline='#0F172A',width=2); d.text((x+w/2,y+23),name,font=bold,fill='#0F172A',anchor='mm')
        for i,f in enumerate(fields): d.text((x+15,y+62+i*32),f,font=small,fill='#111827')
        boxes[name] = (x, y, w, h)
    rels = [("BRANCH","ROOM","1:N"),("ROOM_TYPE","ROOM","1:N"),("CUSTOMER","BOOKING","1:N"),("ROOM","BOOKING","1:N"),("BOOKING","SESSION","0..1:1"),("ROOM","SESSION","1:N"),("SESSION","ORDER","1:N"),("ORDER","ORDER_ITEM","1:N"),("PRODUCT","ORDER_ITEM","1:N"),("SESSION","INVOICE","1:0..1"),("INVOICE","PAYMENT","1:N"),("BRANCH","STAFF","1:N"),("STAFF","ORDER","1:N"),("BRANCH","PRODUCT","1:N")]
    for a,b,label in rels:
        xa,ya,wa,ha=boxes[a]; xb,yb,wb,hb=boxes[b]
        p1=(xa+wa/2,ya+ha/2); p2=(xb+wb/2,yb+hb/2)
        d.line((p1,p2),fill='#64748B',width=2); mx,my=(p1[0]+p2[0])/2,(p1[1]+p2[1])/2; d.rectangle((mx-32,my-12,mx+32,my+12),fill='white'); d.text((mx,my),label,font=small,fill='#B91C1C',anchor='mm')
    d.text((1100,55),'ERD LOGICAL — KARAOKE BOOKING & POS',font=title,fill='#17365D',anchor='mm'); im.save(path)


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr(); el = OxmlElement('w:tblHeader'); el.set(qn('w:val'), 'true'); trPr.append(el)


def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    for i,h in enumerate(headers):
        t.rows[0].cells[i].text = h; shade(t.rows[0].cells[i], "1F4E78")
        for r in t.rows[0].cells[i].paragraphs[0].runs: r.font.color.rgb=RGBColor(255,255,255); r.font.bold=True
    set_repeat_table_header(t.rows[0])
    for row in rows:
        c=t.add_row().cells
        for i,v in enumerate(row): c[i].text=str(v)
    doc.add_paragraph()
    return t


def bullet(doc, text, level=0):
    doc.add_paragraph(text, style="List Bullet" if level == 0 else "List Bullet 2")


def build_doc(path, use_case_img, erd_img):
    doc=Document(); sec=doc.sections[0]
    sec.top_margin=Inches(.7); sec.bottom_margin=Inches(.65); sec.left_margin=Inches(.8); sec.right_margin=Inches(.7)
    styles=doc.styles
    styles['Normal'].font.name='Arial'; styles['Normal'].font.size=Pt(10.5)
    for s,size,color in [('Title',25,'17365D'),('Heading 1',17,'17365D'),('Heading 2',13,'2F5597'),('Heading 3',11,'4472C4')]:
        styles[s].font.name='Arial'; styles[s].font.size=Pt(size); styles[s].font.color.rgb=RGBColor.from_string(color)
    title=doc.add_paragraph(); title.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=title.add_run("TÀI LIỆU ĐẶC TẢ YÊU CẦU PHẦN MỀM\n"); r.bold=True; r.font.size=Pt(24); r.font.color.rgb=RGBColor(23,54,93)
    r=title.add_run("KARAOKE BOOKING & POS"); r.bold=True; r.font.size=Pt(28); r.font.color.rgb=RGBColor(47,85,151)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run("System Design SRS • Phiên bản 1.0\n").bold=True; p.add_run("Môn: AI-Assisted System Design | Ngày: 17/07/2026")
    doc.add_paragraph("\n")
    add_table(doc,["Thuộc tính","Giá trị"],[["Trạng thái","Bản hoàn chỉnh để đánh giá"],["Phạm vi","Chuỗi phòng hát nhiều chi nhánh"],["Chuẩn tham chiếu","IEEE 29148 (cấu trúc rút gọn)"],["Ngôn ngữ","Tiếng Việt"]])
    doc.add_page_break()
    doc.add_heading('MỤC LỤC NỘI DUNG',0)
    for x in ["1. Giới thiệu", "2. Mô tả tổng quan", "3. Yêu cầu chức năng", "4. Quy tắc nghiệp vụ", "5. Mô hình hệ thống", "6. Yêu cầu phi chức năng", "7. Xử lý ngoại lệ", "8. Tiêu chí nghiệm thu & truy vết"]: bullet(doc,x)

    doc.add_heading('1. GIỚI THIỆU',0)
    doc.add_heading('1.1 Mục đích',1)
    doc.add_paragraph('Tài liệu xác định yêu cầu cho nền tảng quản lý chuỗi karaoke, thống nhất giữa chủ đầu tư, quản lý chi nhánh, nhóm phát triển và kiểm thử. Mọi yêu cầu chức năng được gắn mã để có thể truy vết đến tiêu chí nghiệm thu.')
    doc.add_heading('1.2 Phạm vi',1)
    doc.add_paragraph('Hệ thống số hóa vòng đời từ tìm phòng, đặt chỗ, check-in, gọi món, tính tiền đến báo cáo vận hành theo chi nhánh. Phiên bản 1.0 không bao gồm điều khiển thiết bị karaoke, bản quyền nội dung bài hát và kế toán tổng hợp.')
    doc.add_heading('1.3 Thuật ngữ',1)
    add_table(doc,['Thuật ngữ','Định nghĩa'],[['Booking','Yêu cầu giữ một phòng trong khoảng thời gian xác định'],['Session','Phiên sử dụng phòng thực tế từ check-in đến check-out'],['POS','Điểm bán hàng, ghi nhận món và thanh toán'],['No-show','Khách không đến sau thời gian giữ phòng'],['RBAC','Phân quyền theo vai trò'],['SRS','Đặc tả yêu cầu phần mềm']])

    doc.add_heading('2. MÔ TẢ TỔNG QUAN',0)
    doc.add_heading('2.1 Bối cảnh và giả định',1)
    for x in ['Chuỗi có nhiều chi nhánh; mỗi phòng chỉ thuộc một chi nhánh và một loại phòng.', 'Giá phòng phụ thuộc chi nhánh, loại phòng, khung giờ và ngày áp dụng.', 'Đặt phòng được giữ 15 phút khi chờ tiền cọc; quá hạn hệ thống tự giải phóng.', 'Một phiên sử dụng có thể có nhiều order nhưng tối đa một hóa đơn cuối cùng; hóa đơn có thể được thanh toán bằng nhiều phương thức.', 'Múi giờ nghiệp vụ: Asia/Ho_Chi_Minh; tiền tệ: VND; giá lưu bằng số nguyên.']: bullet(doc,x)
    doc.add_heading('2.2 Actors',1)
    add_table(doc,['Actor','Mục tiêu chính','Quyền tiêu biểu'],[
        ['Khách hàng','Tìm và đặt phòng thuận tiện','Xem phòng trống, đặt/hủy, theo dõi booking, gọi món'],
        ['Lễ tân/Thu ngân','Điều phối phòng và thu tiền','Tạo booking, check-in, chuyển phòng, lập hóa đơn, thanh toán'],
        ['Nhân viên phục vụ','Phục vụ món chính xác','Nhận yêu cầu, xác nhận/hoàn tất order'],
        ['Quản lý chi nhánh','Kiểm soát vận hành','Cấu hình phòng/giá/menu, tồn kho, xem báo cáo'],
        ['Quản trị hệ thống','Quản trị nền tảng','Tài khoản, vai trò, chi nhánh, nhật ký']])
    doc.add_heading('2.3 Phân hệ cốt lõi',1)
    add_table(doc,['Module','Trách nhiệm'],[['Booking & Room','Lịch phòng, giữ chỗ, cọc, check-in/out, chuyển phòng'],['POS & Service','Menu, order, trạng thái phục vụ, tồn kho'],['Billing & Payment','Tính tiền phòng/món, giảm giá, thuế, thanh toán, hóa đơn'],['Administration','Chi nhánh, phòng, bảng giá, người dùng, RBAC'],['Reporting','Doanh thu, công suất phòng, sản phẩm bán chạy, đối soát']])
    doc.add_heading('2.4 User Stories ưu tiên',1)
    add_table(doc,['ID','User Story','Ưu tiên'],[
        ['US-01','Là khách hàng, tôi muốn xem phòng trống theo thời gian để chọn đúng phòng.','Must'],['US-02','Là khách hàng, tôi muốn đặt phòng và nộp cọc để được bảo đảm chỗ.','Must'],['US-03','Là lễ tân, tôi muốn check-in đúng booking để tránh cấp trùng phòng.','Must'],['US-04','Là phục vụ, tôi muốn nhận order theo phòng để giao đúng món.','Must'],['US-05','Là thu ngân, tôi muốn hệ thống tổng hợp tiền phòng và món để thanh toán chính xác.','Must'],['US-06','Là quản lý, tôi muốn xem doanh thu và công suất theo ca/ngày.','Should'],['US-07','Là quản trị viên, tôi muốn phân quyền để ngăn truy cập trái phép.','Must']])

    doc.add_heading('3. YÊU CẦU CHỨC NĂNG',0)
    frs=[
        ('FR-01','Tra cứu phòng trống','Khách hàng/Lễ tân','Chi nhánh, thời điểm bắt đầu/kết thúc hợp lệ.','Hệ thống kiểm tra lịch Booking ở trạng thái HOLD/CONFIRMED và Session đang mở; trả phòng phù hợp, sức chứa, giá dự kiến.','Không có phòng: đề xuất khung giờ hoặc loại phòng gần nhất.','Không làm thay đổi dữ liệu.'),
        ('FR-02','Tạo và xác nhận booking','Khách hàng/Lễ tân','Khách đã cung cấp SĐT; phòng còn trống.','Tạo booking HOLD, tính cọc, khóa slot 15 phút; khi cọc thành công chuyển CONFIRMED và gửi mã đặt phòng.','Cọc thất bại/hết hạn: chuyển EXPIRED và giải phóng phòng; yêu cầu trùng lịch: từ chối.','Một booking hợp lệ, không trùng lịch.'),
        ('FR-03','Hủy booking','Khách hàng/Lễ tân','Booking ở HOLD hoặc CONFIRMED; người dùng có quyền.','Kiểm tra chính sách theo thời điểm, tính khoản hoàn, ghi lý do, đổi trạng thái CANCELED, gửi thông báo.','Đã check-in: không được hủy, chuyển quy trình kết thúc phiên.','Slot được giải phóng; giao dịch hoàn được truy vết.'),
        ('FR-04','Check-in và chuyển phòng','Lễ tân','Booking CONFIRMED hoặc khách walk-in; phòng AVAILABLE.','Xác thực mã/SĐT, kiểm tra phòng, tạo Session ACTIVE, đổi phòng OCCUPIED. Khi chuyển phòng, khóa phòng mới rồi cập nhật phiên và giải phóng phòng cũ.','Khách đến sớm: chỉ cho phép trong ngưỡng cấu hình; phòng mới bận: hủy thao tác.','Chỉ một Session ACTIVE trên một phòng.'),
        ('FR-05','Tạo và xử lý order','Khách/Phục vụ','Session ACTIVE; sản phẩm đang bán và đủ tồn.','Tạo order, giữ tồn; phục vụ xác nhận PREPARING, giao món chuyển SERVED; hệ thống trừ tồn chính thức.','Hết món: báo và đề xuất thay thế; hủy trước PREPARING: hoàn tồn giữ.','Order và biến động tồn có nhật ký.'),
        ('FR-06','Checkout và lập hóa đơn','Thu ngân','Session ACTIVE; các order đã SERVED/CANCELED.','Chốt giờ, tính tiền phòng theo bảng giá, cộng món, áp dụng ưu đãi hợp lệ, thuế và tiền cọc; tạo Invoice; đổi Session COMPLETED, phòng CLEANING.','Order chưa xử lý: yêu cầu xử lý; mã giảm giá không hợp lệ: từ chối mã nhưng giữ hóa đơn nháp.','Số tiền bất biến sau khi hóa đơn phát hành.'),
        ('FR-07','Thanh toán đa phương thức','Thu ngân','Invoice ISSUED còn dư nợ.','Chọn tiền mặt/thẻ/QR, ghi từng Payment; khi tổng thành công bằng grand_total, đổi Invoice PAID và phát hành biên nhận.','Gateway timeout: trạng thái PENDING và đối soát, không thử lại mù; trả thừa tiền mặt: tính tiền thối.','Không ghi nhận vượt số dư phải thu.'),
        ('FR-08','Quản lý cấu hình','Quản lý/Quản trị','Đăng nhập và có quyền tương ứng.','CRUD phòng, loại phòng, bảng giá, sản phẩm; kiểm tra hiệu lực và xung đột; ghi audit log.','Không xóa cứng dữ liệu đã phát sinh; chuyển INACTIVE.','Cấu hình mới có version và người thay đổi.'),
        ('FR-09','Báo cáo vận hành','Quản lý','Có quyền trên chi nhánh được giao.','Chọn khoảng thời gian; hệ thống tổng hợp doanh thu đã thanh toán, công suất, hủy/no-show, món bán chạy; cho phép xuất CSV/PDF.','Dữ liệu lớn: tạo tác vụ nền và thông báo khi xong.','Số liệu kèm thời điểm chốt và bộ lọc.'),
        ('FR-10','Quản lý người dùng & RBAC','Quản trị','Tài khoản quản trị hoạt động.','Tạo/khóa tài khoản, gán vai trò và phạm vi chi nhánh; buộc đổi mật khẩu lần đầu; ghi audit.','Không cho quản trị viên tự xóa quyền quản trị cuối cùng.','Quyền mới áp dụng từ phiên đăng nhập kế tiếp.')]
    for fid,name,actor,pre,main,alt,post in frs:
        doc.add_heading(f'{fid} — {name}',1)
        add_table(doc,['Thành phần','Đặc tả'],[['Actor',actor],['Tiền điều kiện',pre],['Luồng chính',main],['Luồng thay thế/ngoại lệ',alt],['Hậu điều kiện',post]])

    doc.add_heading('4. QUY TẮC NGHIỆP VỤ',0)
    add_table(doc,['ID','Quy tắc'],[
        ['BR-01','Hai booking HOLD/CONFIRMED của cùng phòng không được giao nhau theo khoảng [start_at, end_at).'],['BR-02','Booking HOLD tự hết hạn sau 15 phút nếu chưa đủ cọc.'],['BR-03','Tiền phòng được chia theo các đoạn bảng giá có hiệu lực; làm tròn đến 1.000 VND sau tổng.'],['BR-04','Hủy trước ≥ 6 giờ hoàn 100% cọc; từ 2–6 giờ hoàn 50%; dưới 2 giờ hoặc no-show không hoàn. Chính sách phải cấu hình được.'],['BR-05','Giảm giá áp dụng trước thuế; mỗi hóa đơn tối đa một mã khuyến mại, trừ cấu hình cho phép cộng dồn.'],['BR-06','Tồn kho không được âm; giữ tồn khi order tạo và trừ chính thức khi SERVED.'],['BR-07','Chỉ Invoice PAID mới tính vào doanh thu; khoản hoàn được ghi âm ở ngày hoàn.'],['BR-08','Mọi thay đổi giá, giảm giá thủ công, hủy món và hoàn tiền phải lưu actor, thời gian, giá trị trước/sau, lý do.']])
    doc.add_heading('4.1 Máy trạng thái',1)
    add_table(doc,['Đối tượng','Các trạng thái và chuyển đổi hợp lệ'],[['Booking','HOLD → CONFIRMED/EXPIRED/CANCELED; CONFIRMED → CHECKED_IN/CANCELED/NO_SHOW'],['Room','AVAILABLE → RESERVED/OCCUPIED/MAINTENANCE; OCCUPIED → CLEANING → AVAILABLE'],['Order','NEW → CONFIRMED → PREPARING → SERVED; NEW/CONFIRMED → CANCELED'],['Invoice','DRAFT → ISSUED → PARTIALLY_PAID → PAID; ISSUED/PARTIALLY_PAID → VOID (có quyền)'],['Payment','PENDING → SUCCESS/FAILED; SUCCESS → REFUNDED/PARTIALLY_REFUNDED']])

    doc.add_heading('5. MÔ HÌNH HỆ THỐNG',0)
    doc.add_heading('5.1 Sơ đồ Use Case',1); doc.add_picture(str(use_case_img), width=Inches(7.0)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Hình 1 — Tương tác giữa 5 actor và các chức năng cốt lõi.', style='Caption')
    doc.add_heading('5.2 ERD logical',1); doc.add_picture(str(erd_img), width=Inches(7.0)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Hình 2 — Mô hình dữ liệu logical; ký hiệu PK/FK/UQ thể hiện khóa chính, khóa ngoại và duy nhất.', style='Caption')
    doc.add_heading('5.3 Ràng buộc dữ liệu trọng yếu',1)
    for x in ['ROOM: UNIQUE(branch_id, code).', 'BOOKING: start_at < end_at; deposit ≥ 0; chỉ mục (room_id, start_at, end_at, status).', 'ORDER_ITEM: qty > 0; unit_price là ảnh chụp giá tại thời điểm bán.', 'PAYMENT: amount > 0; gateway_transaction_id duy nhất khi có.', 'Dữ liệu tiền dùng DECIMAL(18,0), không dùng kiểu floating point.']: bullet(doc,x)

    doc.add_heading('6. YÊU CẦU PHI CHỨC NĂNG',0)
    add_table(doc,['ID','Nhóm','Chỉ tiêu đo lường'],[
        ['NFR-01','Hiệu năng','P95 tra cứu phòng ≤ 2 giây và P95 thao tác POS ≤ 1 giây với 500 người dùng đồng thời; báo cáo ngày ≤ 10 giây.'],['NFR-02','Sẵn sàng','Uptime tháng ≥ 99,9%, không tính bảo trì báo trước; RTO ≤ 30 phút, RPO ≤ 5 phút.'],['NFR-03','Bảo mật','TLS 1.3 khi truyền; AES-256 khi lưu; mật khẩu Argon2id; khóa sau 5 lần sai/15 phút; MFA bắt buộc với quản trị.'],['NFR-04','Phân quyền','Mọi API kiểm tra RBAC và phạm vi chi nhánh phía server; kiểm thử phải phủ 100% endpoint đặc quyền.'],['NFR-05','Audit','Lưu log nghiệp vụ nhạy cảm tối thiểu 24 tháng; log chống sửa và truy vấn theo actor/thời gian/đối tượng.'],['NFR-06','Khả năng mở rộng','Mở rộng ngang để đạt 3.000 request/phút mà tỷ lệ lỗi < 1% và P95 < 2 giây.'],['NFR-07','Khả dụng','Các luồng đặt phòng và thanh toán hoàn thành trong ≤ 5 bước chính; giao diện responsive từ 360 px.'],['NFR-08','Tương thích','Hỗ trợ hai phiên bản gần nhất của Chrome, Edge, Safari; API JSON UTF-8, version theo /api/v1.'],['NFR-09','Sao lưu','Backup tăng dần mỗi 5 phút, toàn phần hằng ngày; kiểm thử phục hồi hằng quý đạt RTO/RPO.'],['NFR-10','Quan sát','100% request có correlation_id; cảnh báo trong 5 phút khi error rate > 2% trong cửa sổ 10 phút.'],['NFR-11','Riêng tư','Che SĐT trên màn hình không cần thiết; xuất/xóa dữ liệu khách theo chính sách trong ≤ 72 giờ.'],['NFR-12','Toàn vẹn','Các thao tác booking, chuyển phòng, trừ tồn và thanh toán dùng transaction/idempotency key để chống ghi trùng.']])

    doc.add_heading('7. XỬ LÝ NGOẠI LỆ',0)
    add_table(doc,['Mã','Tình huống','Phản ứng hệ thống'],[
        ['EX-01','Hai người cùng đặt phòng','Khóa/kiểm tra xung đột trong transaction; chỉ một booking thành công, bên còn lại nhận gợi ý.'],['EX-02','Gateway thanh toán timeout','Giữ PENDING; truy vấn trạng thái theo idempotency key; không tạo giao dịch mới trước khi đối soát.'],['EX-03','Mất mạng tại POS','Lưu nháp cục bộ tối đa 30 phút; đồng bộ theo mã duy nhất; cảnh báo thao tác cần online như thanh toán.'],['EX-04','Chuyển phòng thất bại giữa chừng','Rollback toàn bộ transaction để phòng cũ vẫn OCCUPIED.'],['EX-05','Sản phẩm vừa hết tồn','Từ chối dòng thiếu, giữ các dòng hợp lệ dưới dạng nháp và đề xuất sản phẩm thay thế.'],['EX-06','Sai lệch tổng tiền','Chặn phát hành hóa đơn nếu tổng chi tiết khác tổng header; ghi log và yêu cầu tính lại.'],['EX-07','Phiên đăng nhập hết hạn','Trả 401, giữ dữ liệu nháp không nhạy cảm; yêu cầu đăng nhập lại.'],['EX-08','Dịch vụ báo cáo quá tải','Đưa vào hàng đợi, trả mã tác vụ và thông báo khi tệp sẵn sàng.']])

    doc.add_heading('8. NGHIỆM THU VÀ TRUY VẾT',0)
    doc.add_heading('8.1 Tiêu chí nghiệm thu mẫu',1)
    add_table(doc,['ID','Given–When–Then'],[
        ['AC-01','Given phòng A có booking 19:00–21:00, When tìm 20:00–22:00, Then phòng A không xuất hiện.'],['AC-02','Given booking HOLD quá 15 phút chưa cọc, When tác vụ hết hạn chạy, Then booking thành EXPIRED và slot được mở.'],['AC-03','Given hai yêu cầu đồng thời cùng phòng/giờ, When xử lý, Then đúng một yêu cầu thành công.'],['AC-04','Given session 120 phút và 3 món đã SERVED, When checkout, Then hóa đơn gồm đủ tiền phòng, món, thuế, cọc và làm tròn đúng BR-03.'],['AC-05','Given thanh toán QR timeout, When người dùng bấm lại, Then hệ thống không thu tiền hai lần.'],['AC-06','Given nhân viên chi nhánh A, When gọi báo cáo chi nhánh B, Then trả 403 và ghi audit.'],['AC-07','Given hủy booking trước 3 giờ, When xác nhận hủy, Then khoản hoàn bằng 50% cọc theo BR-04.'],['AC-08','Given tồn sản phẩm bằng 0, When tạo order, Then bị từ chối và tồn không âm.']])
    doc.add_heading('8.2 Ma trận truy vết',1)
    add_table(doc,['User Story','Yêu cầu','Quy tắc/NFR','Nghiệm thu'],[['US-01','FR-01','BR-01, NFR-01','AC-01'],['US-02','FR-02','BR-01, BR-02, NFR-12','AC-02, AC-03'],['US-03','FR-04','BR-08, NFR-12','AC-03'],['US-04','FR-05','BR-06','AC-08'],['US-05','FR-06, FR-07','BR-03, BR-05, NFR-12','AC-04, AC-05'],['US-06','FR-09','BR-07, NFR-01','AC-06'],['US-07','FR-10','NFR-03, NFR-04, NFR-05','AC-06']])
    doc.add_heading('8.3 Definition of Done',1)
    for x in ['Tất cả yêu cầu Must và các AC liên quan đạt kiểm thử.', 'Không còn lỗi Severity 1–2; kiểm thử tải và phục hồi đạt các NFR.', 'Migration, hướng dẫn vận hành, dashboard/alert và kế hoạch rollback đã được duyệt.', 'Quản lý nghiệp vụ ký nghiệm thu dữ liệu giá, chính sách hủy và mẫu hóa đơn.']: bullet(doc,x)
    doc.add_paragraph('— HẾT TÀI LIỆU —').alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.save(path)


PROMPTS = '''LỊCH SỬ PROMPT — KARAOKE BOOKING & POS\nMục tiêu: thể hiện Iterative Prompting; chỉ lưu prompt, không lưu câu trả lời AI.\n\nPROMPT 01 — KHỞI TẠO VAI TRÒ VÀ MỤC TIÊU\nBạn là Senior IT Consultant kiêm System Analyst, có kinh nghiệm thiết kế hệ thống quản lý chuỗi karaoke. Mục tiêu là đồng hành với tôi xây dựng SRS tiếng Việt theo cấu trúc IEEE 29148 rút gọn. Chưa viết SRS ngay. Trước tiên hãy liệt kê các câu hỏi làm rõ quan trọng về phạm vi, actors, booking, POS, thanh toán và báo cáo; đồng thời chỉ ra các giả định có rủi ro.\n\nPROMPT 02 — CHỐT BỐI CẢNH VÀ GIẢ ĐỊNH\nDựa trên câu hỏi vừa nêu, hãy dùng các quyết định sau: hệ thống phục vụ chuỗi nhiều chi nhánh; có 5 actor gồm Khách hàng, Lễ tân/Thu ngân, Nhân viên phục vụ, Quản lý chi nhánh, Quản trị hệ thống; tiền tệ VND; múi giờ Việt Nam. Phạm vi gồm Booking & Room, POS & Service, Billing & Payment, Administration, Reporting. Ngoài phạm vi: điều khiển thiết bị karaoke, bản quyền bài hát, kế toán tổng hợp. Hãy lập bảng in-scope/out-of-scope, assumptions và 8 rủi ro cần xác nhận. Không tự mở rộng ngoài dữ kiện.\n\nPROMPT 03 — PHÂN TÍCH ACTOR VÀ USER STORY\nHãy phân tích mục tiêu, quyền hạn và pain point của từng actor. Sau đó viết tối thiểu 12 user stories theo mẫu “Là..., tôi muốn..., để...”, gắn mã US-xx và ưu tiên MoSCoW. Kiểm tra để mỗi module có ít nhất 2 user stories và không trao quyền quản trị cho actor nghiệp vụ.\n\nPROMPT 04 — YÊU CẦU CHỨC NĂNG CÓ THỂ KIỂM THỬ\nTừ user stories, tạo danh sách yêu cầu chức năng FR-01... theo từng module. Với mỗi yêu cầu phải có: actor, tiền điều kiện, kích hoạt, luồng chính đánh số, luồng thay thế/ngoại lệ, hậu điều kiện và dữ liệu vào/ra. Dùng động từ “hệ thống phải”; tránh các từ mơ hồ như nhanh, tiện lợi, thân thiện. Chỉ ra user story nào chưa được bao phủ.\n\nPROMPT 05 — REVIEW LOGIC BOOKING VÀ CẠNH TRANH\nĐóng vai reviewer khó tính. Rà soát phần Booking để tìm lỗi trùng lịch, race condition, khách walk-in, đến sớm/muộn, no-show, hủy, cọc thất bại và chuyển phòng. Đề xuất trạng thái Booking/Room/Session cùng các chuyển trạng thái hợp lệ. Sau đó sửa các FR liên quan, bảo đảm hai yêu cầu đồng thời không thể giữ cùng một phòng.\n\nPROMPT 06 — LÀM SÂU POS, TỒN KHO VÀ THANH TOÁN\nRà soát POS và Billing cho các edge case: hết món trong lúc đặt, hủy món đang chế biến, tách nhiều order, nhiều phương thức thanh toán, tiền cọc, hoàn tiền, gateway timeout và người dùng bấm thanh toán lặp. Bổ sung quy tắc tồn kho không âm, idempotency key và trạng thái Payment PENDING/SUCCESS/FAILED/REFUNDED. Cập nhật yêu cầu tương ứng.\n\nPROMPT 07 — QUY TẮC NGHIỆP VỤ\nHãy tổng hợp tối thiểu 8 business rules có mã BR-xx cho xung đột lịch, thời hạn giữ booking, chính sách hoàn cọc, cách tính tiền theo khung giờ, làm tròn VND, giảm giá/thuế, tồn kho và ghi audit. Mỗi quy tắc phải đơn nghĩa và có ví dụ số nếu có phép tính. Nêu rõ quy tắc nào cần cấu hình thay vì hard-code.\n\nPROMPT 08 — USE CASE DIAGRAM\nTạo Mermaid use case diagram (hoặc flowchart mô phỏng use case nếu Mermaid không hỗ trợ cú pháp) cho 5 actor và các chức năng chính. Bảo đảm Quản lý chi nhánh có quản lý phòng/bảng giá/menu/báo cáo; Quản trị hệ thống có người dùng/RBAC; Khách hàng không được thanh toán tại quầy với quyền thu ngân. Sau sơ đồ, lập bảng kiểm để phát hiện actor hoặc use case bị thiếu.\n\nPROMPT 09 — ERD VÀ DATA DICTIONARY\nThiết kế ERD logical cho tối thiểu các thực thể Branch, RoomType, Room, Customer, Booking, Session, Product, Order, OrderItem, Invoice, Payment, Staff, Role. Ghi rõ PK, FK, unique, nullability và quan hệ 1-N/N-N. Xử lý N-N Order–Product qua OrderItem; cho phép một Invoice có nhiều Payment; ngăn booking trùng lịch bằng constraint/transaction phù hợp. Không lưu tiền bằng float.\n\nPROMPT 10 — NFR ĐỊNH LƯỢNG\nViết NFR có mã NFR-xx theo nhóm hiệu năng, sẵn sàng, bảo mật, RBAC, audit, mở rộng, khả dụng, tương thích, backup, quan sát và riêng tư. Mọi NFR phải đo được: dùng P95, số concurrent users, uptime, RTO, RPO, retention, thời gian cảnh báo. Nếu một con số chỉ là giả định, đánh dấu “cần xác nhận với stakeholder”.\n\nPROMPT 11 — EXCEPTION VÀ ACCEPTANCE CRITERIA\nTạo bảng xử lý ngoại lệ cho ít nhất 8 tình huống nguy hiểm. Sau đó viết acceptance criteria dạng Given–When–Then cho các luồng: tìm phòng trống, HOLD hết hạn, đặt đồng thời, checkout, QR timeout, phân quyền chéo chi nhánh, hoàn cọc và hết tồn. Tiêu chí phải có kết quả quan sát được, không chỉ nói “hoạt động đúng”.\n\nPROMPT 12 — REVIEW TRUY VẾT VÀ SỬA THIẾU SÓT\nLập ma trận truy vết User Story → Functional Requirement → Business Rule/NFR → Acceptance Criteria. Đánh dấu mọi yêu cầu mồ côi hoặc user story chưa có kiểm thử. Hãy sửa các khoảng trống trước khi trả bản cuối; không bịa thêm actor hoặc module.\n\nPROMPT 13 — BIÊN TẬP SRS CUỐI\nHợp nhất kết quả thành SRS tiếng Việt chuyên nghiệp gồm: Giới thiệu, Tổng quan, Actors, Modules, User Stories, Functional Requirements, Business Rules, State Models, Use Case, ERD, NFR, Exception Handling, Acceptance Criteria và Traceability Matrix. Dùng thuật ngữ nhất quán, bảng dễ đọc, không để lại lời dẫn hội thoại như “Here is.../Hy vọng hữu ích”. Các sơ đồ phải có chú thích để tôi render thành ảnh khi đưa vào DOCX.\n\nPROMPT 14 — KIỂM ĐỊNH CUỐI THEO RUBRIC\nĐóng vai giảng viên chấm bài theo thang 100: Prompting 30, Logic nghiệp vụ 30, Mermaid/kiến trúc 20, Trình bày 20. Hãy audit bản SRS: (1) đủ ≥3 modules và ≥3 actors, (2) Use Case đúng quyền, (3) ERD đúng cardinality và khóa, (4) NFR đo được, (5) không có lỗi format AI, (6) sơ đồ đã được render thành ảnh. Chỉ trả checklist Pass/Fail và danh sách chỉnh sửa bắt buộc theo mức nghiêm trọng.\n'''


if __name__ == '__main__':
    uc=OUT/'use_case.png'; erd=OUT/'erd.png'
    save_use_case(uc); save_erd(erd)
    build_doc(OUT/'System_Design_SRS.docx', uc, erd)
    (OUT/'History_Prompts.txt').write_text(PROMPTS, encoding='utf-8-sig')
    print(OUT)
